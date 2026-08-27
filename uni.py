import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Configuración de márgenes
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Estilo base
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Portada e Introducción
doc.add_heading('Actividad de Construcción Aplicada (ACA): Matemáticas Especiales', level=0)
p = doc.add_paragraph()
p.add_run('Corporación Unificada Nacional de Educación Superior (CUN)\nDocente: Juan Sebastián Cortés Cruz\nFecha: 2026\n').bold = True

doc.add_heading('1. Introducción', level=1)
doc.add_paragraph(
    'En la arquitectura moderna de sistemas informáticos, las aplicaciones web enfrentan variaciones extremas '
    'en la demanda de tráfico debido a eventos masivos como lanzamientos de productos o ventas relámpago. '
    'Cuando ocurre una ráfaga súbita de peticiones HTTP/HTTPS, el servidor asigna memoria RAM y CPU para encolar '
    'y procesar cada solicitud. Si la tasa de llegada supera la capacidad de drenaje y la memoria se satura, '
    'el sistema operativo activa el OOM Killer, provocando el colapso del servicio.\n\n'
    'El objetivo es formular un modelo dinámico continuo de balance de flujo en la cola de peticiones y '
    'dimensionar la memoria RAM requerida mediante EDOs de 1.er orden, Transformada de Laplace e implementación computacional.'
)

# Sección 2
doc.add_heading('2. Planteamiento, Modelación y Resolución del Problema', level=1)
doc.add_heading('Contexto y Pregunta Central', level=2)
doc.add_paragraph(
    '¿Cuál es la función temporal q(t) que describe el número de peticiones acumuladas en el buffer del servidor '
    'tras un pico repentino de tráfico, cuánta memoria RAM pico requerirá la infraestructura para no colapsar '
    'y en qué tiempo el sistema alcanza el 98% de su estado estacionario?'
)

doc.add_heading('Variables y Parámetros', level=2)
table_data = [
    ['Símbolo', 'Tipo', 'Descripción', 'Unidades', 'Valor'],
    ['t', 'Var. indep.', 'Tiempo desde inicio del pico', 's', 't ≥ 0'],
    ['q(t)', 'Var. dep.', 'Longitud de cola / peticiones', 'req', 'Variable'],
    ['λ(t)', 'Parámetro', 'Tasa de llegada de peticiones', 'req/s', 'λ(t) = 500 · u(t)'],
    ['α', 'Parámetro', 'Coeficiente de servicio del servidor', 's⁻¹', '0.5 s⁻¹'],
    ['q₀', 'Cond. inicial', 'Peticiones en t = 0⁻', 'req', '0 req'],
    ['m_r', 'Parámetro', 'Memoria consumida por petición', 'MB/req', '4.0 MB/req'],
    ['M_base', 'Parámetro', 'Memoria estática del SO y runtime', 'MB', '1024 MB (1 GB)'],
    ['M(t)', 'Var. dep.', 'Memoria RAM total consumida', 'MB', 'Variable'],
    ['M_máx', 'Restricción', 'Capacidad total de RAM física', 'MB', '8192 MB (8 GB)']
]

table = doc.add_table(rows=len(table_data), cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for r_idx, row in enumerate(table_data):
    for c_idx, val in enumerate(row):
        cell = table.cell(r_idx, c_idx)
        cell.text = val
        if r_idx == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True

doc.add_heading('Resolución Matemática Paso a Paso', level=2)

pasos = [
    ("Paso 1: Formulación de la Ecuación Diferencial (EDO)",
     "Por conservación de flujo continuo:\n"
     "dq(t)/dt = λ(t) - α · q(t)\n"
     "dq(t)/dt + α · q(t) = λ₀ · u(t),  con q(0) = 0"),

    ("Paso 2: Aplicación de la Transformada de Laplace",
     "Aplicando ℒ{·} a ambos lados con ℒ{dq/dt} = s·Q(s) - q(0) y ℒ{u(t)} = 1/s:\n"
     "s·Q(s) - 0 + α·Q(s) = λ₀ / s\n"
     "Q(s)·(s + α) = λ₀ / s\n"
     "Q(s) = λ₀ / [ s · (s + α) ]"),

    ("Paso 3: Descomposición en Fracciones Parciales",
     "λ₀ / [ s · (s + α) ] = A / s + B / (s + α)\n"
     "λ₀ = A·(s + α) + B·s\n"
     "• Para s = 0: λ₀ = A·α  =>  A = λ₀ / α\n"
     "• Para s = -α: λ₀ = -B·α  =>  B = -λ₀ / α\n"
     "Q(s) = (λ₀ / α) · (1 / s) - (λ₀ / α) · [ 1 / (s + α) ]"),

    ("Paso 4: Transformada Inversa de Laplace",
     "Aplicando ℒ⁻¹{·}:\n"
     "q(t) = (λ₀ / α) · [ 1 - e^(-α·t) ]\n"
     "Con valores numéricos (λ₀ = 500 req/s, α = 0.5 s⁻¹):\n"
     "q(t) = (500 / 0.5) · [ 1 - e^(-0.5·t) ] = 1000 · [ 1 - e^(-0.5·t) ]  [req]"),

    ("Paso 5: Modelo Dinámico de Memoria RAM",
     "M(t) = M_base + m_r · q(t)\n"
     "M(t) = 1024 + 4.0 · 1000 · [ 1 - e^(-0.5·t) ] = 1024 + 4000 · [ 1 - e^(-0.5·t) ]  [MB]")
]

for title, content in pasos:
    doc.add_heading(title, level=3)
    p = doc.add_paragraph(content)
    p.paragraph_format.left_indent = Inches(0.2)

doc.add_heading('Verificación de Resultados', level=2)
doc.add_paragraph(
    "1. Sustitución en EDO: dq/dt = 500·e^(-0.5t). Sustituyendo: 500·e^(-0.5t) + 0.5(1000 - 1000·e^(-0.5t)) = 500 = λ₀ (Cumple).\n"
    "2. Teorema Valor Inicial: lím(s->∞) s·Q(s) = lím(s->∞) 500/(s+0.5) = 0 req (Cumple q(0)=0).\n"
    "3. Teorema Valor Final: lím(s->0) s·Q(s) = lím(s->0) 500/(s+0.5) = 1000 req (Cumple).\n"
    "4. Tiempo al 98%: t_s = 4τ = 4(1/α) = 4(2s) = 8.0 s."
)

# Sección 3
doc.add_heading('3. Resultados, Análisis e Interpretación', level=1)
res_data = [
    ['Métrica / Parámetro', 'Fórmula', 'Valor', 'Unidades'],
    ['Constante de tiempo (τ)', '1 / α', '2.00', 's'],
    ['Tiempo de estabilización (98%)', '4τ', '8.00', 's'],
    ['Longitud de cola estacionaria (q_ss)', 'λ₀ / α', '1000', 'req'],
    ['Memoria RAM Pico (M_peak)', 'M_base + m_r · q_ss', '5024', 'MB (4.91 GB)'],
    ['RAM física instalada', 'M_máx', '8192', 'MB (8.00 GB)'],
    ['Margen de seguridad operacional', 'M_máx - M_peak', '3168', 'MB (3.09 GB)']
]

res_table = doc.add_table(rows=len(res_data), cols=4)
res_table.style = 'Table Grid'
res_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for r_idx, row in enumerate(res_data):
    for c_idx, val in enumerate(row):
        cell = res_table.cell(r_idx, c_idx)
        cell.text = val
        if r_idx == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True

doc.add_paragraph(
    '\nInterpretación Técnica: El servidor no colapsa ante el pico de 500 req/s, estabilizándose a los 8 segundos '
    'en 5024 MB de consumo total, manteniendo un margen seguro de 3.09 GB (38.7% de memoria libre).'
)

doc.save('ACA_Matematicas_Especiales.docx')
print("Archivo 'ACA_Matematicas_Especiales.docx' generado exitosamente.")